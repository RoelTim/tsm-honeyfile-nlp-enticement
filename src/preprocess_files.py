"""
Preprocess Deception and Local Context Files with Spacy
* remove stop words, punctuation, numbers
* apply lemmatisation
* apply named-entity recognition (NER)
"""

import os
import glob
import json
import docx
import spacy

class EntityRetokenizeComponent:
    """ merges tokens that are entities """
    def __init__(self, nlp):
        pass
    def __call__(self, doc):
        with doc.retokenize() as retokenizer:
            for ent in doc.ents:
                retokenizer.merge(doc[ent.start:ent.end],
                                  attrs={"LEMMA": str(doc[ent.start:ent.end])})
        return doc

def get_category(file_path):
    """
    define category of file based on file_path
    """
    if 'theater' in file_path:
        category = 'theater'
    elif 'plants' in file_path:
        category = 'plants'
    elif 'computer' in file_path:
        category = 'computer'
    elif 'customs' in file_path or 'abfnotices' in file_path:
        category = 'customs'
    else:
        category = 'unknown'
    return category

def nlp_preprocess_steps(text):
    """
    remove stop words, punctuation, numbers
    apply lemmatisation
    apply named-entity recognition (NER)
    """
    nlp = spacy.load("en_core_web_sm")
    nlp.add_pipe(EntityRetokenizeComponent(nlp))
    text = text.lower()
    tokens = nlp(text)
    processed_text = []
    for sen in tokens.sents:
        sen = str(sen)
        nlp_sen = nlp(sen)
        sent = []
        for token in nlp_sen:
            if token.pos_ != 'PUNCT' and token.pos_ != 'NUM':
                if not token.is_stop and len(token) > 2 and '\n' not in token.lemma_:
                    sent.append(str(token.lemma_))
        processed_text.append(sent)
    return processed_text

def preprocess_deception_files():
    """ preprocess deception files """
    print('start preprocessing deception files')
    path_this_file = os.path.dirname(os.path.abspath(__file__))
    path_dec_files = os.path.abspath(os.path.join(path_this_file, '..', 'data/raw/honeyfiles'))
    dec_files = glob.glob(path_dec_files + '/**/*.docx', recursive=True)
    decep_files = {}
    for dec_file in dec_files:
        doc = docx.Document(dec_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        decep_files[dec_file] = '\n'.join(full_text)

    doc2vec_text = {}
    lda_text = {}
    sbmtm_text = {}
    preprocessed_text = {}

    for count, (i, j) in enumerate(decep_files.items()):
        if count % 100 == 0:
            print(count, ' of ', len(decep_files))
        processed_text = nlp_preprocess_steps(j)
        sbmtm = [item for sublist in processed_text for item in sublist]
        doc2vec_text[i] = ' '.join(sbmtm)
        lda_text[i] = processed_text
        sbmtm_text[i] = processed_text
        category = get_category(j)
        metadata = i + '...' + category + '...N/A...dec...doc2vec'
        preprocessed_text[metadata] = ' '.join(sbmtm)
        metadata = i + '...' + category + '...N/A...dec...lda'
        preprocessed_text[metadata] = processed_text
    output_file = os.path.abspath(os.path.join(path_this_file, \
                                               '..', \
                                               'data/processed/preprocessed_text.txt'))
    json.dump(preprocessed_text, open(output_file, 'w'))
    print('finished preprocessing deception files')

def preproces_local_context_files():
    """ preprocess local context files """
    print('start preprocessing local context files')
    path_this_file = os.path.dirname(os.path.abspath(__file__))
    path_local_files = os.path.abspath(os.path.join(path_this_file, '..', 'data/raw/local_context'))
    local_files = glob.glob(path_local_files + '/**/*.docx', recursive=True)
    print(len(local_files))
    output_file = os.path.abspath(os.path.join(path_this_file, \
                                               '..', \
                                               'data/processed/preprocessed_text.txt'))

    with open(output_file) as file:
        preprocessed_text = json.load(file)

    for count, i in enumerate(local_files):
        if count % 100 == 0:
            print(count, ' of ', len(local_files))
        try:
            doc = docx.Document(i)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            j = '\n'.join(full_text)
            processed_text = nlp_preprocess_steps(j)
            sbmtm = [item for sublist in processed_text for item in sublist]
            category = get_category(i)
            metadata = i + '...' + category + '...1...local...doc2vec'
            preprocessed_text[metadata] = ' '.join(sbmtm)
            metadata = i + '...' + category + '...1...local...lda'
            preprocessed_text[metadata] = processed_text
        except:
            print('skip: ', i)

    output_file = os.path.abspath(os.path.join(path_this_file, \
                                               '..', \
                                               'data/processed/preprocessed_text.txt'))
    json.dump(preprocessed_text, open(output_file, 'w'))
    print('finished preprocessing local context files')

if __name__ == "__main__":
    preprocess_deception_files()
    preproces_local_context_files()
